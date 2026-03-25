# ============================================================
# utils/file_handler.py
# Detects file types and extracts searchable content from
# images, PDFs, text files, Office documents, and plain text.
# ============================================================

from __future__ import annotations

import io
import logging
import os
import re
from pathlib import Path
from typing import Any

import chardet

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional heavy imports – gracefully degrade if a library is not installed
# ---------------------------------------------------------------------------
try:
    from PIL import Image

    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("Pillow not installed – image handling will be limited.")

try:
    import imagehash

    IMAGEHASH_AVAILABLE = True
except ImportError:
    IMAGEHASH_AVAILABLE = False
    logger.warning("imagehash not installed – perceptual hashing disabled.")

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not installed – PDF text extraction disabled.")

try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx  # python-docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed – .docx extraction disabled.")

try:
    import openpyxl

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not installed – .xlsx extraction disabled.")

try:
    from striprtf.striprtf import rtf_to_text

    STRIPRTF_AVAILABLE = True
except ImportError:
    STRIPRTF_AVAILABLE = False
    logger.warning("striprtf not installed – .rtf extraction disabled.")

# ============================================================
# CONSTANTS
# ============================================================

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif"}
PDF_EXTENSIONS = {".pdf"}
TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".xml",
    ".html",
    ".htm",
    ".yaml",
    ".yml",
    ".ini",
    ".cfg",
    ".log",
    ".rst",
}
DOC_EXTENSIONS = {".doc", ".docx", ".odt", ".xls", ".xlsx", ".ppt", ".pptx", ".rtf"}

ALL_SUPPORTED = IMAGE_EXTENSIONS | PDF_EXTENSIONS | TEXT_EXTENSIONS | DOC_EXTENSIONS


# ============================================================
# PUBLIC API
# ============================================================


class FileInfo:
    """
    Container for all extracted information about an uploaded file or
    a plain-text input.

    Attributes
    ----------
    category : str
        One of 'image', 'pdf', 'text', 'doc', 'plain_text', 'unknown'.
    text_content : str
        Best-effort plain-text representation of the file's content.
    image : PIL.Image.Image | None
        Loaded image object (only for image files).
    image_hash : imagehash.ImageHash | None
        Perceptual hash of the image (only for image files).
    image_phash : imagehash.ImageHash | None
        PHash variant for the image.
    image_dhash : imagehash.ImageHash | None
        DHash variant for the image.
    image_whash : imagehash.ImageHash | None
        Wavelet hash variant for the image.
    metadata : dict
        Additional metadata (size, page count, encoding, …).
    filename : str
        Original filename or '<text_input>'.
    filepath : Path | None
        Path to the file on disk (None for plain-text inputs).
    error : str | None
        Description of any error that occurred during extraction.
    """

    def __init__(self) -> None:
        self.category: str = "unknown"
        self.text_content: str = ""
        self.image: Any = None  # PIL.Image
        self.image_hash: Any = None  # imagehash.ImageHash (ahash)
        self.image_phash: Any = None  # imagehash.ImageHash (phash)
        self.image_dhash: Any = None  # imagehash.ImageHash (dhash)
        self.image_whash: Any = None  # imagehash.ImageHash (whash)
        self.metadata: dict = {}
        self.filename: str = "<unknown>"
        self.filepath: Path | None = None
        self.error: str | None = None

    def to_dict(self) -> dict:
        """Serialisable summary (no heavy objects)."""
        return {
            "filename": self.filename,
            "filepath": str(self.filepath) if self.filepath else None,
            "category": self.category,
            "text_content_preview": self.text_content[:500]
            if self.text_content
            else "",
            "text_length": len(self.text_content),
            "has_image": self.image is not None,
            "has_image_hash": self.image_hash is not None,
            "metadata": self.metadata,
            "error": self.error,
        }


# -----------------------------------------------------------------------
# Entry points
# -----------------------------------------------------------------------


def process_uploaded_file(filepath: str | Path, original_filename: str) -> FileInfo:
    """
    Process an uploaded file and return a populated FileInfo object.

    Parameters
    ----------
    filepath : str | Path
        Path to the file saved on disk.
    original_filename : str
        The original name of the uploaded file (used for extension detection).
    """
    filepath = Path(filepath)
    info = FileInfo()
    info.filepath = filepath
    info.filename = original_filename
    info.metadata["file_size_bytes"] = (
        filepath.stat().st_size if filepath.exists() else 0
    )

    ext = Path(original_filename).suffix.lower()
    info.category = _categorise(ext)

    try:
        if info.category == "image":
            _process_image(filepath, info)
        elif info.category == "pdf":
            _process_pdf(filepath, info)
        elif info.category == "text":
            _process_text_file(filepath, info)
        elif info.category == "doc":
            _process_doc_file(filepath, info, ext)
        else:
            # Try to read as text anyway
            _process_text_file(filepath, info)
            info.category = "unknown"
    except Exception as exc:
        logger.exception("Error processing uploaded file %s", filepath)
        info.error = str(exc)

    return info


def process_plain_text(text: str) -> FileInfo:
    """
    Wrap a raw text string in a FileInfo for searching.

    Parameters
    ----------
    text : str
        The plain text entered by the user.
    """
    info = FileInfo()
    info.category = "plain_text"
    info.filename = "<text_input>"
    info.filepath = None
    info.text_content = _clean_text(text)
    info.metadata["char_count"] = len(info.text_content)
    info.metadata["word_count"] = len(info.text_content.split())
    return info


def process_file_for_indexing(filepath: str | Path) -> FileInfo:
    """
    Process an arbitrary file found on disk (local search indexing).
    Uses only the file extension to determine type.
    """
    filepath = Path(filepath)
    info = FileInfo()
    info.filepath = filepath
    info.filename = filepath.name

    try:
        info.metadata["file_size_bytes"] = filepath.stat().st_size
    except OSError:
        info.metadata["file_size_bytes"] = 0

    ext = filepath.suffix.lower()
    info.category = _categorise(ext)

    try:
        if info.category == "image":
            _process_image(filepath, info)
        elif info.category == "pdf":
            _process_pdf(filepath, info)
        elif info.category == "text":
            _process_text_file(filepath, info)
        elif info.category == "doc":
            _process_doc_file(filepath, info, ext)
    except Exception as exc:
        logger.debug("Could not fully process %s: %s", filepath, exc)
        info.error = str(exc)

    return info


# ============================================================
# INTERNAL PROCESSORS
# ============================================================


def _categorise(ext: str) -> str:
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in PDF_EXTENSIONS:
        return "pdf"
    if ext in TEXT_EXTENSIONS:
        return "text"
    if ext in DOC_EXTENSIONS:
        return "doc"
    return "unknown"


# -----------------------------------------------------------------------
# Image processing
# -----------------------------------------------------------------------


def _process_image(filepath: Path, info: FileInfo) -> None:
    if not PIL_AVAILABLE:
        info.error = "Pillow is not installed."
        return

    img = Image.open(filepath)
    img = img.convert("RGB")  # normalise mode
    info.image = img
    info.metadata["image_mode"] = img.mode
    info.metadata["image_size"] = img.size  # (width, height)
    info.metadata["image_format"] = img.format

    if IMAGEHASH_AVAILABLE:
        try:
            info.image_hash = imagehash.average_hash(img)
            info.image_phash = imagehash.phash(img)
            info.image_dhash = imagehash.dhash(img)
            info.image_whash = imagehash.whash(img)
        except Exception as exc:
            logger.warning("Could not compute image hashes for %s: %s", filepath, exc)

    # Attempt OCR text extraction if pytesseract is available
    info.text_content = _ocr_image(img)


def _ocr_image(img: Any) -> str:
    """Try to extract text from an image via Tesseract OCR."""
    try:
        import pytesseract  # optional – not in requirements by default

        text = pytesseract.image_to_string(img)
        return _clean_text(text)
    except ImportError:
        return ""
    except Exception as exc:
        logger.debug("OCR failed: %s", exc)
        return ""


# -----------------------------------------------------------------------
# PDF processing
# -----------------------------------------------------------------------


def _process_pdf(filepath: Path, info: FileInfo) -> None:
    text_parts: list[str] = []

    # --- Primary: pdfplumber (better accuracy) ---
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(filepath) as pdf:
                info.metadata["page_count"] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
        except Exception as exc:
            logger.warning("pdfplumber failed on %s: %s", filepath, exc)

    # --- Fallback: PyPDF2 ---
    if not text_parts and PYPDF2_AVAILABLE:
        try:
            with open(filepath, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                info.metadata["page_count"] = len(reader.pages)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("PyPDF2 failed on %s: %s", filepath, exc)

    combined = "\n".join(text_parts)
    info.text_content = _clean_text(combined)
    info.metadata["extracted_chars"] = len(info.text_content)


# -----------------------------------------------------------------------
# Plain text / markup processing
# -----------------------------------------------------------------------


def _process_text_file(filepath: Path, info: FileInfo) -> None:
    raw_bytes = _safe_read_bytes(filepath)
    if raw_bytes is None:
        info.error = f"Could not read file: {filepath}"
        return

    encoding = _detect_encoding(raw_bytes)
    info.metadata["encoding"] = encoding

    try:
        text = raw_bytes.decode(encoding, errors="replace")
    except (LookupError, UnicodeDecodeError):
        text = raw_bytes.decode("utf-8", errors="replace")

    ext = filepath.suffix.lower()
    if ext in {".html", ".htm"}:
        text = _strip_html(text)
    elif ext == ".json":
        text = _flatten_json(text)
    elif ext in {".xml"}:
        text = _strip_xml_tags(text)

    info.text_content = _clean_text(text)
    info.metadata["extracted_chars"] = len(info.text_content)


# -----------------------------------------------------------------------
# Office document processing
# -----------------------------------------------------------------------


def _process_doc_file(filepath: Path, info: FileInfo, ext: str) -> None:
    if ext == ".docx" and DOCX_AVAILABLE:
        _extract_docx(filepath, info)
    elif ext in {".xlsx", ".xls"} and OPENPYXL_AVAILABLE:
        _extract_xlsx(filepath, info)
    elif ext == ".rtf" and STRIPRTF_AVAILABLE:
        _extract_rtf(filepath, info)
    else:
        # Last-ditch: try reading as text
        _process_text_file(filepath, info)


def _extract_docx(filepath: Path, info: FileInfo) -> None:
    try:
        document = docx.Document(str(filepath))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        # Also grab table cell text
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        info.text_content = _clean_text("\n".join(paragraphs))
        info.metadata["paragraph_count"] = len(document.paragraphs)
    except Exception as exc:
        logger.warning("docx extraction failed for %s: %s", filepath, exc)
        info.error = str(exc)


def _extract_xlsx(filepath: Path, info: FileInfo) -> None:
    try:
        wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    parts.append(row_text)
        info.text_content = _clean_text("\n".join(parts))
        info.metadata["sheet_count"] = len(wb.worksheets)
        wb.close()
    except Exception as exc:
        logger.warning("xlsx extraction failed for %s: %s", filepath, exc)
        info.error = str(exc)


def _extract_rtf(filepath: Path, info: FileInfo) -> None:
    try:
        raw_bytes = _safe_read_bytes(filepath)
        if raw_bytes:
            encoding = _detect_encoding(raw_bytes)
            rtf_str = raw_bytes.decode(encoding, errors="replace")
            plain = rtf_to_text(rtf_str)
            info.text_content = _clean_text(plain)
    except Exception as exc:
        logger.warning("rtf extraction failed for %s: %s", filepath, exc)
        info.error = str(exc)


# ============================================================
# SIMILARITY HELPERS  (used by searchers)
# ============================================================


def compute_image_hashes(image_input: str | Path | Any) -> dict[str, Any]:
    """
    Compute all perceptual hashes for an image given a filepath or a PIL Image.

    Returns a dict with keys: 'ahash', 'phash', 'dhash', 'whash'.
    Returns an empty dict if imagehash / Pillow is not available.
    """
    if not PIL_AVAILABLE or not IMAGEHASH_AVAILABLE:
        return {}

    try:
        if isinstance(image_input, (str, Path)):
            img = Image.open(image_input).convert("RGB")
        else:
            img = image_input.convert("RGB")

        return {
            "ahash": imagehash.average_hash(img),
            "phash": imagehash.phash(img),
            "dhash": imagehash.dhash(img),
            "whash": imagehash.whash(img),
        }
    except Exception as exc:
        logger.debug("Hash computation failed: %s", exc)
        return {}


def image_similarity_score(hashes_a: dict, hashes_b: dict) -> float:
    """
    Return a similarity score in [0.0, 1.0] between two hash dicts.
    1.0 means identical, 0.0 means completely different.
    """
    if not hashes_a or not hashes_b:
        return 0.0

    scores: list[float] = []
    for key in ("ahash", "phash", "dhash", "whash"):
        ha = hashes_a.get(key)
        hb = hashes_b.get(key)
        if ha is not None and hb is not None:
            # Max hash bit length is 64 for 8x8 hash
            max_bits = 64
            distance = ha - hb
            scores.append(1.0 - distance / max_bits)

    if not scores:
        return 0.0
    return max(0.0, min(1.0, sum(scores) / len(scores)))


def extract_text_from_file(filepath: str | Path) -> str:
    """
    Convenience function: return just the extracted text from any file.
    Returns an empty string on failure.
    """
    info = process_file_for_indexing(filepath)
    return info.text_content


# ============================================================
# UTILITY HELPERS
# ============================================================


def _safe_read_bytes(filepath: Path, max_bytes: int = 50 * 1024 * 1024) -> bytes | None:
    """Read up to max_bytes from a file. Returns None on error."""
    try:
        size = filepath.stat().st_size
        with open(filepath, "rb") as fh:
            return fh.read(min(size, max_bytes))
    except OSError as exc:
        logger.debug("Cannot read %s: %s", filepath, exc)
        return None


def _detect_encoding(raw: bytes) -> str:
    """Detect the character encoding of raw bytes."""
    result = chardet.detect(raw[:10_000])  # sample first 10 KB
    return result.get("encoding") or "utf-8"


def _clean_text(text: str) -> str:
    """Normalise whitespace and remove non-printable characters."""
    # Remove null bytes and other control characters except newlines/tabs
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    # Collapse multiple spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse more than 2 consecutive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_html(html: str) -> str:
    """Remove HTML tags and decode common entities."""
    # Remove <script> and <style> blocks entirely
    html = re.sub(
        r"<(script|style)[^>]*>.*?</(script|style)>",
        " ",
        html,
        flags=re.DOTALL | re.IGNORECASE,
    )
    # Remove all remaining tags
    html = re.sub(r"<[^>]+>", " ", html)
    # Decode common entities
    entities = {
        "&amp;": "&",
        "&lt;": "<",
        "&gt;": ">",
        "&nbsp;": " ",
        "&quot;": '"',
        "&#39;": "'",
    }
    for entity, char in entities.items():
        html = html.replace(entity, char)
    return html


def _strip_xml_tags(xml: str) -> str:
    """Remove XML tags, leaving only text nodes."""
    return re.sub(r"<[^>]+>", " ", xml)


def _flatten_json(json_str: str) -> str:
    """
    Extract all string values from a JSON document as a flat text blob.
    Falls back to returning the raw string if parsing fails.
    """
    import json

    try:
        data = json.loads(json_str)
        parts: list[str] = []
        _collect_strings(data, parts)
        return " ".join(parts)
    except json.JSONDecodeError:
        return json_str


def _collect_strings(obj: Any, out: list[str]) -> None:
    """Recursively collect all string values from a JSON-parsed object."""
    if isinstance(obj, str):
        out.append(obj)
    elif isinstance(obj, dict):
        for v in obj.values():
            _collect_strings(v, out)
    elif isinstance(obj, (list, tuple)):
        for item in obj:
            _collect_strings(item, out)


# ============================================================
# FILENAME / METADATA HELPERS
# ============================================================


def get_file_extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def is_image(filename: str) -> bool:
    return get_file_extension(filename) in IMAGE_EXTENSIONS


def is_pdf(filename: str) -> bool:
    return get_file_extension(filename) in PDF_EXTENSIONS


def is_text_file(filename: str) -> bool:
    return get_file_extension(filename) in TEXT_EXTENSIONS


def is_doc_file(filename: str) -> bool:
    return get_file_extension(filename) in DOC_EXTENSIONS


def is_supported(filename: str) -> bool:
    return get_file_extension(filename) in ALL_SUPPORTED


def human_readable_size(num_bytes: int) -> str:
    """Convert a byte count into a human-readable string."""
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if abs(num_bytes) < 1024.0:
            return f"{num_bytes:3.1f} {unit}"
        num_bytes /= 1024.0
    return f"{num_bytes:.1f} PB"
