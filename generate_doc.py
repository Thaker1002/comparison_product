from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(30, 80, 160)):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        pass
    else:
        run.text = text
    run.font.color.rgb = RGBColor(*color)
    return p

def h1(text): return heading(text, 1, (15, 60, 140))
def h2(text): return heading(text, 2, (30, 100, 180))
def h3(text): return heading(text, 3, (50, 130, 200))

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F4F8')
    p._p.pPr.append(shading) if p._p.pPr is not None else None
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 100)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def kv(key, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(30, 80, 160)
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("\nPriceHunt — Multi-Country Product Comparison App")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(15, 60, 140)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("Full Technical & Feature Documentation")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(80, 80, 80)

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run(f"\nGenerated: {datetime.datetime.now().strftime('%B %d, %Y')}\n")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(120, 120, 120)

divider()
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. App Overview")
para(
    "PriceHunt is a full-stack web application that lets users compare product prices "
    "across multiple online marketplaces in different countries. It also includes a ride "
    "fare estimator powered by Google Maps. The app supports multi-language display, GPS-based "
    "location detection, user authentication, usage tracking, and an admin dashboard.",
    size=11
)

doc.add_paragraph()
h2("1.1 Key Features")
features = [
    "Product price search across Amazon, Shopee, Lazada, Noon, Flipkart, and more",
    "Image-based product search (upload a photo to find matching products)",
    "Ride fare estimator with interactive Google Maps route picker",
    "GPS auto-detection to pre-select country and language",
    "Multi-language UI: English, Thai, Bahasa Indonesia, Tagalog, Hindi, Tamil, Arabic, Français",
    "User registration (auto-generated password sent by email)",
    "User login with JWT authentication — persisted across browser refresh",
    "Forgot/reset password flow",
    "Every search is logged to SQLite (usage tracking)",
    "Admin dashboard — user list, search stats, per-user usage history",
    "Progressive Web App (PWA) with installable manifest + service worker",
    "Fully responsive — works on desktop, tablet, and mobile",
    "Dark-themed modern UI built with Tailwind CSS + shadcn/ui",
]
for f in features:
    bullet(f)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. TECH STACK
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. Technology Stack")

h2("2.1 Frontend")
stack_fe = [
    ("Framework", "React 18 with TypeScript"),
    ("Build Tool", "Vite 5"),
    ("Styling", "Tailwind CSS v3 + shadcn/ui component library"),
    ("Maps", "@react-google-maps/api — Google Maps JavaScript API"),
    ("i18n", "react-i18next — 8 locales (en, th, id, tl, hi, ta, ar, fr)"),
    ("HTTP", "axios"),
    ("State", "React useState / useEffect hooks (no Redux)"),
    ("PWA", "Custom service worker (sw.js) + Web App Manifest"),
]
for k, v in stack_fe:
    kv(k, v)

doc.add_paragraph()
h2("2.2 Backend")
stack_be = [
    ("Runtime", "Node.js with ES Modules"),
    ("Framework", "Express 4"),
    ("Database", "SQLite via better-sqlite3"),
    ("Auth", "bcryptjs (password hashing) + jsonwebtoken (JWT)"),
    ("Email", "nodemailer (SMTP)"),
    ("Search API", "Tavily Search API"),
    ("Translation", "OpenRouter → Claude Sonnet 4"),
    ("Image Search", "Cloudinary + Tavily image endpoint"),
    ("Config", "dotenv"),
]
for k, v in stack_be:
    kv(k, v)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. PROJECT STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Project Structure")
para("The repository is organized into three top-level folders:")
doc.add_paragraph()
code_block(
"""product_comparison/
├── backend/
│   ├── package.json          — Node dependencies
│   └── src/
│       ├── index.js          — Main Express server (entry point)
│       ├── auth.js           — Auth, user, admin routes
│       ├── db.js             — SQLite database init & schema
│       ├── countries.js      — Country config (marketplaces per country)
│       ├── url-validation.js — URL sanitisation helpers
│       └── url-validation.test.js
│
├── frontend/
│   ├── package.json
│   ├── vite.config.ts        — Vite config + /api proxy to :3001
│   ├── tailwind.config.js
│   ├── index.html            — PWA entry point
│   └── src/
│       ├── App.tsx           — Root component
│       ├── main.tsx          — React DOM entry / i18n bootstrap
│       ├── index.css
│       ├── i18n/             — Translation JSON files per locale
│       ├── components/
│       │   ├── SearchBar.tsx
│       │   ├── ResultsGrid.tsx
│       │   ├── ProductCard.tsx
│       │   ├── TaxiTab.tsx
│       │   ├── Header.tsx
│       │   ├── ImageUpload.tsx
│       │   ├── AuthModal.tsx
│       │   └── AdminDashboard.tsx
│       ├── lib/
│       │   ├── api.ts        — Frontend API calls (search, image search)
│       │   ├── url-validation.ts
│       │   └── utils.ts
│       └── types/
│           └── index.ts      — Shared TypeScript interfaces
│
└── mobile/                   — Stub for future React Native app
"""
)
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. FRONTEND — COMPONENT BY COMPONENT
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. Frontend Components")

# 4.1 App.tsx
h2("4.1  App.tsx  —  Root Component")
para(
    "App.tsx is the top-level component. It owns all global state and orchestrates "
    "every feature of the application.",
    size=11
)
doc.add_paragraph()
h3("State variables")
state_vars = [
    ("activeCategory", "\"purchase\" | \"ride\" | \"flights\" — selected tab"),
    ("country", "ISO-2 country code: TH, ID, PH, MY, SG, IN, AE, US, CA"),
    ("language", "Locale code: en, th, id, tl, hi, ta, ar, fr"),
    ("geoStatus", "\"idle\" | \"loading\" | \"done\" | \"error\" — GPS detection progress"),
    ("query", "Current text-search query string"),
    ("authUser", "Logged-in user object (null when not authenticated)"),
    ("authToken", "JWT string stored in localStorage"),
    ("showAdmin", "Boolean — toggles AdminDashboard overlay"),
    ("results / loading / error", "Product search results state"),
    ("filters", "Active price/rating/sort filters (DEFAULT_FILTERS)"),
]
for k, v in state_vars:
    bullet(f"{k}  →  {v}")

doc.add_paragraph()
h3("Key functions in App.tsx")
fns = [
    ("handleAuth(user, token)", "Called after login/register. Saves user + token to state AND localStorage (persists on page refresh)."),
    ("handleLogout()", "Clears authUser, authToken, and removes both localStorage keys."),
    ("trackUsage(event_type, query, results_count)", "POSTs to /api/auth/track with Bearer token. Called on every search so activity is recorded."),
    ("GPS useEffect", "Runs once on mount. navigator.geolocation.getCurrentPosition → reverse-geocodes with Nominatim (OpenStreetMap) → auto-sets country + language."),
    ("handleSearch(q, f)", "Calls searchProducts() from lib/api.ts, then calls trackUsage(). Updates results state."),
    ("handleImageSearch(file)", "Uploads image to /api/image-search, updates results."),
]
for fn, desc in fns:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(f"{fn}  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

doc.add_paragraph()
h3("Country → Language mapping (auto-select on GPS)")
code_block(
"TH → th   |   ID → id   |   PH → tl   |   MY → en\n"
"SG → en   |   IN → hi   |   AE → ar   |   US → en   |   CA → en"
)

doc.add_paragraph()
h3("Supported Countries")
countries = [
    ("TH", "Thailand", "🇹🇭"),
    ("ID", "Indonesia", "🇮🇩"),
    ("PH", "Philippines", "🇵🇭"),
    ("MY", "Malaysia", "🇲🇾"),
    ("SG", "Singapore", "🇸🇬"),
    ("IN", "India", "🇮🇳"),
    ("AE", "UAE / Dubai", "🇦🇪"),
    ("US", "USA", "🇺🇸"),
    ("CA", "Canada", "🇨🇦"),
]
for code, name, _ in countries:
    bullet(f"{code}  —  {name}")

doc.add_paragraph()

# 4.2 AuthModal
h2("4.2  AuthModal.tsx  —  Authentication Gate")
para(
    "Shown as a full-screen overlay before any app content. Users must register or log in "
    "before they can use the app.",
    size=11
)
doc.add_paragraph()
h3("Tabs")
bullet("register  —  New user sign-up form")
bullet("login     —  Returning user sign-in form")
bullet("forgot    —  Password reset by email")

doc.add_paragraph()
h3("Register flow")
steps = [
    "User fills: Full Name, Email, Mobile number, Notification preference (Email / WhatsApp / SMS)",
    "POST /api/auth/register → backend generates a secure random password, hashes it, stores in SQLite",
    "Response includes the plain-text password — shown on screen for 4 seconds",
    "Email is sent to user with their credentials (if SMTP is configured)",
    "After 4 seconds, onAuth(user, token) is called — user is logged in automatically",
]
for s in steps:
    bullet(s)

doc.add_paragraph()
h3("Login flow")
steps = [
    "User enters Email + Password",
    "POST /api/auth/login → backend verifies password hash, returns JWT + user object",
    "onAuth() saves to state + localStorage",
]
for s in steps:
    bullet(s)

doc.add_paragraph()
h3("Forgot password flow")
steps = [
    "User enters their email address",
    "POST /api/auth/forgot-password → backend generates a new random password, updates hash in DB",
    "New password is sent to user's email",
    "New password is also shown on screen",
]
for s in steps:
    bullet(s)

doc.add_paragraph()

# 4.3 SearchBar
h2("4.3  SearchBar.tsx  —  Search Input")
para("Text search bar with voice input toggle. Supports keyboard Enter to trigger search.", size=11)
bullet("Voice search: Web Speech API (SpeechRecognition) — fills the input with recognised text")
bullet("Submit triggers handleSearch() in App.tsx")
bullet("Displays real-time loading spinner during search")

doc.add_paragraph()

# 4.4 ResultsGrid
h2("4.4  ResultsGrid.tsx  —  Product Results Layout")
para(
    "Renders the product cards returned by the search API in a responsive grid. "
    "Includes filter controls (price range, min rating, sort order).",
    size=11
)
bullet("Filters: min price, max price, minimum star rating, sort by price/rating/relevance")
bullet("Grouped view by marketplace — collapse/expand each group")
bullet("Empty-state message when no results found")
bullet("Loading skeleton cards while search is in progress")

doc.add_paragraph()

# 4.5 ProductCard
h2("4.5  ProductCard.tsx  —  Individual Product Card")
para("Displays a single product search result.", size=11)
fields = [
    "Product image (with fallback placeholder)",
    "Product title",
    "Price with currency symbol",
    "Star rating badge",
    "Review count",
    "Marketplace name + badge",
    "\"View Deal\" button — opens product URL in new tab",
    "Translate button — calls /api/translate to translate description into current language",
]
for f in fields:
    bullet(f)

doc.add_paragraph()

# 4.6 TaxiTab
h2("4.6  TaxiTab.tsx  —  Ride Fare Estimator")
para(
    "Powered by Google Maps JavaScript API. Estimates ride costs between two points "
    "in supported countries.",
    size=11
)
doc.add_paragraph()
h3("How it works")
steps = [
    "User clicks two points on the Google Map (Point A and Point B)",
    "Or types addresses in the from/to text fields with autocomplete",
    "Distance and duration are calculated via the Directions API",
    "Fare estimates are shown for: Grab, Bolt, Gojek, InDrive (depending on country)",
    "Fare = base fare + (per-km rate × distance) + (per-min rate × duration)",
]
for s in steps:
    bullet(s)

doc.add_paragraph()
h3("Map implementation")
bullet("useLoadScript({ googleMapsApiKey }) — loads Google Maps SDK once")
bullet("GoogleMap component renders the interactive map")
bullet("Marker components for Point A (green) and Point B (red)")
bullet("Polyline draws the route between the two markers")

doc.add_paragraph()

# 4.7 ImageUpload
h2("4.7  ImageUpload.tsx  —  Image-Based Search")
para(
    "Users can take a photo or upload an image from their device. "
    "The image is sent to the backend, which extracts product keywords "
    "and runs a marketplace search.",
    size=11
)
bullet("Accepts: JPEG, PNG, WEBP, max 10 MB")
bullet("Drag-and-drop or click-to-open file picker")
bullet("Preview thumbnail shown before submission")
bullet("POST multipart/form-data to /api/image-search")

doc.add_paragraph()

# 4.8 AdminDashboard
h2("4.8  AdminDashboard.tsx  —  Admin Panel")
para(
    "Full-screen overlay accessible only to users with is_admin = 1 in the database. "
    "Admin must enter the ADMIN_SECRET to get an admin JWT.",
    size=11
)
doc.add_paragraph()
h3("Admin login")
bullet("Text field for ADMIN_SECRET → POST /api/admin/login → stores admin JWT in localStorage")

h3("Statistics cards")
cards = [
    "Total registered users",
    "Total searches (all time)",
    "Today's searches",
    "Average searches per user",
]
for c in cards:
    bullet(c)

h3("Charts / lists")
bullet("Top Search Queries (most searched terms + count)")
bullet("Searches by Country (country code + count)")

h3("Users table")
cols = ["Name", "Email", "Mobile", "Notify via", "Total searches", "Date joined", "Last login", "View button"]
for c in cols:
    bullet(c)

h3("User detail panel")
bullet("Click 'View' on any user → loads full usage event history")
bullet("Shows: event type, query text, country, results count, timestamp for every event")

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. BACKEND
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Backend")
para("The backend is a Node.js / Express server. All routes use ES Module syntax.", size=11)
doc.add_paragraph()

h2("5.1  index.js  —  Main Server")
kv("Port", "3001 (or process.env.PORT)")
kv("Entry point", "node src/index.js  (start command in package.json)")
doc.add_paragraph()
h3("Registered routes")
routes = [
    ("POST", "/api/search",           "Product search — calls Tavily, filters/cleans results"),
    ("POST", "/api/image-search",     "Upload image → Cloudinary → extract keywords → Tavily search"),
    ("POST", "/api/translate",        "Translate text via OpenRouter (Claude Sonnet 4)"),
    ("GET",  "/api/countries",        "List all supported countries + their marketplace configs"),
    ("GET",  "/api/countries/:code",  "Get single country config"),
    ("GET",  "/api/stats",            "In-memory telemetry: total searches, products returned, uptime"),
    ("POST", "/api/auth/register",    "Register new user"),
    ("POST", "/api/auth/login",       "Login existing user"),
    ("POST", "/api/auth/forgot-password", "Reset password by email"),
    ("POST", "/api/auth/track",       "Log a usage event (requires Bearer JWT)"),
    ("POST", "/api/admin/login",      "Get admin JWT using ADMIN_SECRET"),
    ("GET",  "/api/admin/users",      "List all users + search stats (requires admin JWT)"),
    ("GET",  "/api/admin/usage/:userId", "Usage events for one user (requires admin JWT)"),
    ("GET",  "/api/admin/stats",      "Aggregate search statistics (requires admin JWT)"),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
for cell, text, w in zip(hdr, ["Method", "Path", "Description"], [1.0, 2.4, 3.5]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.width = Inches(w)

for method, path, desc in routes:
    row = tbl.add_row().cells
    row[0].text = method
    _g = 140 if method == "GET" else 0
    _r = 0 if method == "GET" else 180
    row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(_r, _g, 0)
    row[0].paragraphs[0].runs[0].font.size = Pt(9.5)
    row[1].text = path
    row[1].paragraphs[0].runs[0].font.name = "Courier New"
    row[1].paragraphs[0].runs[0].font.size = Pt(9)
    row[2].text = desc
    row[2].paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph()

h2("5.2  auth.js  —  Authentication Module")
doc.add_paragraph()
h3("Helper functions")
fns = [
    ("generatePassword(len=10)", "Generates a cryptographically random password from a safe character set (no ambiguous chars like 0/O/1/l)."),
    ("sendWelcomeEmail(user, plainPassword)", "Sends credentials email via nodemailer. Silently skips if SMTP is not configured in .env."),
    ("sendResetEmail(user, newPassword)", "Sends new password to user's email after a reset request."),
    ("requireAuth(req,res,next)", "Express middleware — verifies JWT Bearer token in Authorization header."),
    ("requireAdmin(req,res,next)", "Express middleware — verifies JWT AND checks payload.isAdmin === true."),
]
for fn, desc in fns:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(f"{fn}  ")
    r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

doc.add_paragraph()
h3("POST /api/auth/register")
steps = [
    "Validates name, email, mobile are present",
    "Checks email is not already taken (UNIQUE constraint)",
    "Calls generatePassword(10) to create a random password",
    "Hashes password with bcrypt (salt rounds = 10)",
    "Inserts user row into SQLite users table",
    "Calls sendWelcomeEmail() — sends credentials to user's email",
    "Returns: { user, token, password } — password only shown once",
]
for s in steps: bullet(s)

doc.add_paragraph()
h3("POST /api/auth/login")
steps = [
    "Finds user by email",
    "bcrypt.compare(plainPassword, storedHash)",
    "Updates last_login timestamp",
    "Returns JWT (signed with JWT_SECRET, expires in 30 days) + user object",
]
for s in steps: bullet(s)

doc.add_paragraph()
h3("POST /api/auth/forgot-password")
steps = [
    "Finds user by email (does not reveal if email exists — sends 200 either way)",
    "Generates new random password, hashes it, updates users table",
    "Calls sendResetEmail() to email new password",
    "Also returns new password in JSON response so user sees it on screen",
]
for s in steps: bullet(s)

doc.add_paragraph()
h3("POST /api/auth/track  (requires JWT)")
steps = [
    "requireAuth middleware verifies the token",
    "Inserts row into usage_events: user_id, event_type, query, country, results_count, timestamp",
]
for s in steps: bullet(s)

doc.add_paragraph()

h2("5.3  db.js  —  Database")
para("Initialises a SQLite database at backend/data/users.db using better-sqlite3.", size=11)
doc.add_paragraph()
h3("Schema")
code_block(
"""-- users table
CREATE TABLE IF NOT EXISTS users (
  id          INTEGER PRIMARY KEY AUTOINCREMENT,
  name        TEXT    NOT NULL,
  email       TEXT    NOT NULL UNIQUE,
  mobile      TEXT    NOT NULL,
  password    TEXT    NOT NULL,       -- bcrypt hash
  notify_via  TEXT    NOT NULL DEFAULT 'email',  -- 'email' | 'whatsapp' | 'sms'
  is_verified INTEGER NOT NULL DEFAULT 1,
  is_admin    INTEGER NOT NULL DEFAULT 0,
  created_at  TEXT    NOT NULL DEFAULT (datetime('now')),
  last_login  TEXT
);

-- usage_events table
CREATE TABLE IF NOT EXISTS usage_events (
  id            INTEGER PRIMARY KEY AUTOINCREMENT,
  user_id       INTEGER NOT NULL REFERENCES users(id),
  event_type    TEXT    NOT NULL,  -- 'search' | 'image_search' | 'ride_estimate'
  query         TEXT,
  country       TEXT,
  results_count INTEGER DEFAULT 0,
  created_at    TEXT    NOT NULL DEFAULT (datetime('now'))
);"""
)
doc.add_paragraph()
h3("Settings")
bullet("WAL journal mode enabled for better concurrent read performance")
bullet("data/ directory is auto-created if it does not exist")
bullet("users.db is listed in .gitignore — never committed to version control")

doc.add_paragraph()

h2("5.4  countries.js  —  Country Configuration")
para(
    "Exports a COUNTRIES map that associates each country code with: "
    "display name, currency, currency symbol, and the list of marketplace "
    "domains to search (e.g. amazon.com, shopee.co.th, lazada.co.th).",
    size=11
)
doc.add_paragraph()

h2("5.5  url-validation.js  —  URL Sanitisation")
para(
    "Utility functions used by the search endpoint to clean, normalise and validate "
    "product URLs returned by Tavily before sending them to the frontend.",
    size=11
)
helpers = [
    "normaliseUrl(url) — removes tracking params, enforces HTTPS",
    "isValidProductUrl(url) — checks domain against allowed marketplace list",
    "isValidImageUrl(url) — checks URL looks like a real product image",
    "buildFallbackUrl(marketplace, query) — constructs a safe search URL if no valid URL found",
    "isPlausibleShopeeUrl(url) — Shopee-specific heuristic check",
    "BLOCKED_DOMAINS — set of ad/tracking domains to reject",
    "FAKE_PATH_PATTERNS — regex list for obviously fake/placeholder paths",
]
for h in helpers:
    bullet(h)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. DATABASE
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Database")
para("SQLite is used for simplicity and zero-infrastructure cost. The database file lives at:", size=11)
code_block("backend/data/users.db")
doc.add_paragraph()
para("On Render.com (production), a persistent disk is needed or the DB resets on each deploy. Alternatively, replace better-sqlite3 with a hosted Postgres.", size=11)
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. AUTHENTICATION & SECURITY
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Authentication & Security")

h2("7.1 User Authentication")
bullet("Passwords are hashed with bcrypt (cost factor 10) — never stored in plain text")
bullet("JWT tokens expire after 30 days")
bullet("Tokens are stored in browser localStorage")
bullet("All protected routes verify the JWT in requireAuth middleware")
bullet("Admin routes require a separate JWT issued only when ADMIN_SECRET is presented")

doc.add_paragraph()
h2("7.2 Admin Access")
steps = [
    "Admin clicks the Admin button (visible only if user.isAdmin === true)",
    "Opens AdminDashboard overlay",
    "Enters ADMIN_SECRET in the login form",
    "POST /api/admin/login returns an admin JWT",
    "All admin API calls include this JWT as Bearer token",
    "Admin JWT is stored in localStorage as adminToken",
]
for s in steps:
    bullet(s)

doc.add_paragraph()
h2("7.3 Password Security")
bullet("Auto-generated passwords use only unambiguous characters to avoid confusion")
bullet("Character set: ABCDEFGHJKLMNPQRSTUVWXYZ + abcdefghjkmnpqrstuvwxyz + 23456789@#!")
bullet("No 0/O/1/l/I characters that look similar")
bullet("Passwords are communicated via email (SMTP) and shown on screen once")

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. ENVIRONMENT VARIABLES
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Environment Variables")
para("Create a file called .env in the backend/ folder with the following variables:", size=11)
code_block(
"""# Required for search
TAVILY_API_KEY=your_tavily_api_key_here

# Required for translation
OPENROUTER_API_KEY=your_openrouter_api_key_here

# Required for authentication
JWT_SECRET=your_very_long_random_secret_here
ADMIN_SECRET=your_admin_password_here

# Optional: image search
CLOUDINARY_CLOUD_NAME=your_cloud_name
CLOUDINARY_API_KEY=your_api_key
CLOUDINARY_API_SECRET=your_api_secret

# Optional: email delivery
SMTP_HOST=smtp.gmail.com
SMTP_PORT=587
SMTP_USER=your_email@gmail.com
SMTP_PASS=your_app_password
SMTP_FROM=PriceHunt <your_email@gmail.com>

# Optional: server port (default 3001)
PORT=3001"""
)
doc.add_paragraph()
para("If SMTP variables are missing, the app still works — emails are silently skipped and passwords are shown on screen only.", size=11)
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. HOW TO RUN LOCALLY
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. How to Run Locally")

h2("9.1 Prerequisites")
bullet("Node.js v18 or later")
bullet("npm v9 or later")
bullet("Python 3 (for Machine Search module only)")

doc.add_paragraph()
h2("9.2 Install dependencies")
code_block(
"""# Backend
cd backend
npm install

# Frontend
cd ../frontend
npm install"""
)

doc.add_paragraph()
h2("9.3 Configure environment")
code_block(
"""# Copy the example and fill in your keys
cd backend
copy .env.example .env
# Then edit .env with your API keys"""
)

doc.add_paragraph()
h2("9.4 Start the servers")
code_block(
"""# Terminal 1 — Backend (port 3001)
cd backend
npm start

# Terminal 2 — Frontend (port 5173)
cd frontend
npm run dev"""
)
para("Or use the provided start.bat / start.ps1 to launch both together.", size=11)
doc.add_paragraph()
h2("9.5 Access on phone / tablet (same WiFi)")
code_block(
"""# Start Vite with --host to expose on local network
cd frontend
npm run dev -- --host

# Vite will print your local IP, e.g.:
#   http://10.217.6.169:5173/
# Open that URL on your phone"""
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. DEPLOYMENT (Render.com)
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. Deployment on Render.com")
para("The project includes a render.yaml defining two services:", size=11)

h2("10.1 Backend Web Service")
bullet("Build command: cd backend && npm install")
bullet("Start command: node src/index.js")
bullet("Add all .env variables as Render environment variables")
bullet("Add a Render Persistent Disk mounted at /opt/render/project/src/data to keep users.db")

doc.add_paragraph()
h2("10.2 Frontend Static Site")
bullet("Build command: cd frontend && npm install && npm run build")
bullet("Publish directory: frontend/dist")
bullet("Add environment variable: VITE_API_URL = https://your-backend-name.onrender.com")

doc.add_paragraph()
h2("10.3 CORS")
para("The backend currently allows all origins (cors()). For production, restrict to your Render frontend URL:", size=11)
code_block('app.use(cors({ origin: "https://your-frontend.onrender.com" }));')
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. PWA
# ═══════════════════════════════════════════════════════════════════════════════
h1("11. Progressive Web App (PWA)")
para("The frontend is installable as a PWA on Android and iOS.", size=11)
bullet("frontend/public/manifest.json — defines app name, icons, theme color, display mode")
bullet("frontend/public/sw.js — service worker for caching and offline support")
bullet("On Chrome/Edge: browser shows 'Add to Home Screen' prompt automatically")
bullet("App icon set included in frontend/public/icons/")
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. I18N / LOCALIZATION
# ═══════════════════════════════════════════════════════════════════════════════
h1("12. Multi-Language Support (i18n)")
para("Built with react-i18next. Translation files are in frontend/src/i18n/.", size=11)
langs = [
    ("en", "English", "Default fallback language"),
    ("th", "Thai (ไทย)", "Auto-selected for Thailand"),
    ("id", "Bahasa Indonesia", "Auto-selected for Indonesia"),
    ("tl", "Tagalog", "Auto-selected for Philippines"),
    ("hi", "Hindi (हिन्दी)", "Auto-selected for India"),
    ("ta", "Tamil (தமிழ்)", "Available for Indian users"),
    ("ar", "Arabic (العربية)", "Auto-selected for UAE, RTL layout"),
    ("fr", "Français", "French"),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for cell, text in zip(hdr, ["Code", "Language", "Notes"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
for code, name, note in langs:
    row = tbl.add_row().cells
    row[0].text = code
    row[0].paragraphs[0].runs[0].font.size = Pt(10)
    row[1].text = name
    row[1].paragraphs[0].runs[0].font.size = Pt(10)
    row[2].text = note
    row[2].paragraphs[0].runs[0].font.size = Pt(10)
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  13. GITHUB
# ═══════════════════════════════════════════════════════════════════════════════
h1("13. GitHub Repository")
kv("URL",    "https://github.com/Thaker1002/comparison_product")
kv("Branch", "main")
kv("Latest commit", "715c501 — Add forgot/reset password flow")
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  14. GOOGLE MAPS API KEY
# ═══════════════════════════════════════════════════════════════════════════════
h1("14. Google Maps API Key")
kv("Key",   "AIzaSyAhlQd0NyVC-ex0G5ySR91US0GMcXwy2NQ")
kv("Used in", "TaxiTab.tsx (Maps JavaScript API)")
kv("Restricted in", "Google Cloud Console — restrict to your domain in production")
para("APIs enabled: Maps JavaScript API, Directions API, Places API, Geocoding API", size=11)
divider()

# ═══════════════════════════════════════════════════════════════════════════════
#  15. KNOWN LIMITATIONS & FUTURE IMPROVEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
h1("15. Known Limitations & Future Improvements")

h2("Current limitations")
limits = [
    "SQLite is not suitable for high-traffic production — migrate to PostgreSQL for scale",
    "WhatsApp / SMS password delivery via Twilio is not yet implemented (stored in notify_via but only email works)",
    "Flights tab shows 'Coming Soon' — not yet implemented",
    "Product URLs from Tavily are occasionally broken — URL sanitisation catches most but not all",
    "Rate limiting is not implemented — production deployments should add express-rate-limit",
    "The database file resets on Render.com free tier deploys unless a persistent disk is attached",
]
for l in limits:
    bullet(l)

doc.add_paragraph()
h2("Suggested future improvements")
improvements = [
    "Twilio integration for WhatsApp and SMS password delivery",
    "Change password page for logged-in users",
    "User profile page (edit name, mobile, notification preference)",
    "Price history charts using Chart.js or Recharts",
    "Email verification on registration (currently is_verified=1 by default)",
    "Rate limiting and brute-force protection on /api/auth/login",
    "Flights tab using Skyscanner or Amadeus API",
    "PostgreSQL migration for production scalability",
    "Redis caching for repeated search queries",
    "React Native mobile app (mobile/ folder is already scaffolded)",
]
for i in improvements:
    bullet(i)

divider()

# ─── Save ─────────────────────────────────────────────────────────────────────
OUTPUT_PATH = r"c:\Users\thake\Downloads\ZED\product_comparision\PriceHunt_Documentation.docx"
doc.save(OUTPUT_PATH)
print(f"✅ Word document saved to:\n{OUTPUT_PATH}")
